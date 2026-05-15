from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "FBR-Frontend-Requirements-Report.md"
OUTPUT = ROOT / "FBR-Frontend-Requirements-Report.docx"

BURGUNDY = "5A0F24"
DARK_BURGUNDY = "3B0A18"
ROSE = "F7E8ED"
ROSE_2 = "FDF7F8"
GOLD = "C99A3E"
CHARCOAL = "222222"
MUTED = "666666"
BORDER = "D9C7CE"
CODE_BG = "F6F1F3"


def set_shading(element, fill):
    tc_pr = element._tc.get_or_add_tcPr() if hasattr(element, "_tc") else element._p.get_or_add_pPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")


def add_paragraph_border(paragraph, color=GOLD, size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def strip_md_markers(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text


def add_inline(paragraph, text, font_size=None, color=None):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            apply_run_style(run, font_size, color)
        token = match.group(0)
        content = token[2:-2] if token.startswith("**") else token[1:-1]
        run = paragraph.add_run(content)
        apply_run_style(run, font_size, color)
        if token.startswith("**"):
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BURGUNDY if color is None else color)
        elif token.startswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10 if font_size is None else font_size)
            run.font.color.rgb = RGBColor.from_string(DARK_BURGUNDY)
        elif token.startswith("*"):
            run.italic = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        apply_run_style(run, font_size, color)


def apply_run_style(run, font_size=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    if font_size is not None:
        run.font.size = Pt(font_size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def make_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(CHARCOAL)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(DARK_BURGUNDY)
    styles["Title"].paragraph_format.space_after = Pt(4)

    styles["Heading 1"].font.size = Pt(15.5)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor.from_string(BURGUNDY)
    styles["Heading 1"].paragraph_format.space_before = Pt(15)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)

    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor.from_string(BURGUNDY)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)

    styles["Heading 3"].font.size = Pt(11)
    styles["Heading 3"].font.bold = True
    styles["Heading 3"].font.color.rgb = RGBColor.from_string(CHARCOAL)
    styles["Heading 3"].paragraph_format.space_before = Pt(7)
    styles["Heading 3"].paragraph_format.space_after = Pt(3)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(8.7)
        code.font.color.rgb = RGBColor.from_string(DARK_BURGUNDY)
        code.paragraph_format.line_spacing = 1.0
        code.paragraph_format.space_after = Pt(0)

    if "Report Small" not in styles:
        small = styles.add_style("Report Small", WD_STYLE_TYPE.PARAGRAPH)
        small.font.name = "Arial"
        small._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        small.font.size = Pt(8)
        small.font.color.rgb = RGBColor.from_string(MUTED)


def setup_page(doc, title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run(title)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BURGUNDY)
    add_paragraph_border(hp, color=BORDER, size="6")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run("Page ")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(fp)


def parse_table(lines):
    rows = []
    for line in lines:
        text = line.strip()
        if text.startswith("|"):
            text = text[1:]
        if text.endswith("|"):
            text = text[:-1]
        rows.append([cell.strip() for cell in text.split("|")])
    if len(rows) >= 2 and all(re.match(r"^:?-{3,}:?$", c.strip()) for c in rows[1]):
        rows.pop(1)
    return rows


def column_widths(rows, total=10160):
    cols = max(len(row) for row in rows)
    weights = [1] * cols
    for i in range(cols):
        values = [strip_md_markers(row[i]) for row in rows if i < len(row)]
        avg = sum(min(max(len(v), 7), 55) for v in values) / max(len(values), 1)
        weights[i] = max(8, avg)
    s = sum(weights)
    widths = [int(total * w / s) for w in weights]
    widths[-1] += total - sum(widths)
    return widths


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = column_widths(rows)
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        tr_pr = table.rows[r_idx]._tr.get_or_add_trPr()
        if r_idx == 0:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell, color=BORDER)
            if r_idx == 0:
                set_shading(cell, BURGUNDY)
            elif r_idx % 2 == 0:
                set_shading(cell, ROSE_2)
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            paragraph.paragraph_format.space_after = Pt(0)
            if r_idx == 0:
                add_inline(paragraph, strip_md_markers(text), font_size=8.2, color="FFFFFF")
                for run in paragraph.runs:
                    run.bold = True
            else:
                add_inline(paragraph, text, font_size=8.2)
                if len(strip_md_markers(text)) <= 18 and "\n" not in text:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [10160])
    cell = table.cell(0, 0)
    set_shading(cell, CODE_BG)
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    set_cell_borders(cell, color="E4D7DC")
    paragraph = cell.paragraphs[0]
    paragraph.style = "Code Block"
    paragraph.text = ""
    for i, line in enumerate(lines):
        if i:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(DARK_BURGUNDY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_metadata_table(doc, metadata):
    table = doc.add_table(rows=len(metadata), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, [1600, 8560])
    for idx, (label, value) in enumerate(metadata):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        for cell in (left, right):
            set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
            set_cell_borders(cell, color="E3CED5")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_shading(left, BURGUNDY)
        set_shading(right, ROSE)
        lp = left.paragraphs[0]
        lp.text = ""
        add_inline(lp, label, font_size=8.5, color="FFFFFF")
        for run in lp.runs:
            run.bold = True
        rp = right.paragraphs[0]
        rp.text = ""
        add_inline(rp, value, font_size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def convert():
    text = SOURCE.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    make_styles(doc)
    title = strip_md_markers(lines[0].lstrip("# ").strip())
    setup_page(doc, title)

    i = 0
    pending_metadata = []
    started_body = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(7)
            add_paragraph_border(p, color=GOLD, size="12")
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            add_code_block(doc, block)
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue

        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            if pending_metadata:
                add_metadata_table(doc, pending_metadata)
                pending_metadata = []
            level = len(heading.group(1))
            text_value = heading.group(2).strip()
            if level == 1:
                p = doc.add_paragraph(style="Title")
                add_inline(p, text_value)
                add_paragraph_border(p, color=GOLD, size="12")
            else:
                p = doc.add_paragraph(style=f"Heading {level - 1}")
                add_inline(p, text_value)
                if level == 2:
                    add_paragraph_border(p, color="E7D6DC", size="4")
            started_body = True
            i += 1
            continue

        meta = re.match(r"^\*\*([^:]+):\s*(.*?)\*\*$", stripped)
        if meta and started_body and not any(p.text.startswith("1. Project Overview") for p in doc.paragraphs):
            pending_metadata.append((meta.group(1), meta.group(2)))
            i += 1
            continue

        if pending_metadata:
            add_metadata_table(doc, pending_metadata)
            pending_metadata = []

        bullet = re.match(r"^-\s+(.*)$", stripped)
        number = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet or number:
            content = (bullet or number).group(1)
            style = "List Bullet" if bullet else "List Number"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.14)
            p.paragraph_format.space_after = Pt(5)
            add_inline(p, content)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_inline(p, stripped)
        i += 1

    if pending_metadata:
        add_metadata_table(doc, pending_metadata)

    # Avoid a conspicuous blank last page when the source ends with a rule + closing line.
    for section in doc.sections:
        section.start_type = WD_SECTION_START.CONTINUOUS

    doc.core_properties.title = title
    doc.core_properties.subject = "Frontend requirements report"
    doc.core_properties.author = "Techionik"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    convert()
